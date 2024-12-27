import os
import re
import json
import traceback
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

# Carregar variáveis de ambiente do arquivo .env
load_dotenv()

def create_docx_from_json(arquivo_json, arquivo_saida='curriculo.docx'):
    """Cria um documento Word formatado a partir de dados de um currículo em JSON."""
    # Carregar dados JSON
    with open(arquivo_json, 'r', encoding='utf-8') as f:
        dados = json.load(f)

    # Criar um novo Documento
    doc = Document()

    # Definir fonte padrão
    estilo = doc.styles['Normal']
    estilo.font.name = 'Calibri'
    estilo.font.size = Pt(11)
    estilo.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor preta

    # Função para adicionar espaço entre seções
    def adicionar_espaco():
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Nome (Cabeçalho Centralizado)
    informacoes_pessoais = dados.get('informacoes_pessoais', {})
    nome = informacoes_pessoais.get('nome', 'Nome Não Encontrado')
    paragrafo_nome = doc.add_paragraph(nome)
    nome_run = paragrafo_nome.runs[0]
    nome_run.bold = True
    nome_run.font.size = Pt(16)
    paragrafo_nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Informações de Contato (Alinhado à esquerda)
    adicionar_espaco()
    contato = f"""Cidade: {informacoes_pessoais.get('cidade', 'N/A')}
                    Bairro: {informacoes_pessoais.get('bairro', 'N/A')}
                    Email: {informacoes_pessoais.get('email', 'N/A')} 
                    Telefone: {informacoes_pessoais.get('telefone', 'N/A')}
                    Posição: {informacoes_pessoais.get('cargo', 'N/A')}"""
    paragrafo_contato = doc.add_paragraph(contato)
    paragrafo_contato.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Formação
    adicionar_espaco()
    paragrafo_formacao = doc.add_paragraph('Formação:')
    paragrafo_formacao.runs[0].bold = True
    paragrafo_formacao.runs[0].font.size = Pt(12)
    
    educacao = dados.get('educacao', [])
    for item in educacao:
        instituicao = item.get('instituicao', 'Instituição Não Especificada')
        grau = item.get('grau', 'Grau Não Especificado')
        ano_formatura = item.get('ano_formatura', 'Ano Não Especificado')
        
        paragrafo_educacao = doc.add_paragraph(f' • {grau}, {instituicao} - finalizado em {ano_formatura}')
        paragrafo_educacao.paragraph_format.left_indent = Pt(18)

    # Certificações
    adicionar_espaco()
    paragrafo_certificacoes = doc.add_paragraph('Certificações:')
    paragrafo_certificacoes.runs[0].bold = True
    paragrafo_certificacoes.runs[0].font.size = Pt(12)
    
    certificacoes = dados.get('certificacoes', [])
    for item in certificacoes:
        certificado = item.get('certificado', 'Certificado Não Especificado')
        paragrafo_certificacao = doc.add_paragraph(f' • {certificado}')
        paragrafo_certificacao.paragraph_format.left_indent = Pt(18)

    # Resumo de Qualificações
    adicionar_espaco()
    paragrafo_qualificacoes = doc.add_paragraph('Resumo de Qualificações:')
    paragrafo_qualificacoes.runs[0].bold = True
    paragrafo_qualificacoes.runs[0].font.size = Pt(12)
    
    qualificacoes = dados.get('resumo_qualificacoes', [])
    if qualificacoes:
        primeira_qualificacao = qualificacoes[0]
        resumo = primeira_qualificacao.get('resumo', '')
        if resumo:
            doc.add_paragraph(resumo)
        
        qualificacoes_chave = primeira_qualificacao.get('qualificacoes_chave', [])
        for qual in qualificacoes_chave:
            texto_qualificacao = qual.get('qualificacao', '') if isinstance(qual, dict) else str(qual)
            paragrafo_qual = doc.add_paragraph(f' • {texto_qualificacao}')
            paragrafo_qual.paragraph_format.left_indent = Pt(18)

    # Experiência Profissional
    adicionar_espaco()
    paragrafo_experiencia = doc.add_paragraph('Experiências:')
    paragrafo_experiencia.runs[0].bold = True
    paragrafo_experiencia.runs[0].font.size = Pt(12)
    
    experiencias = dados.get('experiencia_profissional', [])
    for trabalho in experiencias:
        adicionar_espaco()
        empresa = trabalho.get('empresa', 'Empresa Não Especificada')
        cargo = trabalho.get('cargo', 'Cargo Não Especificado')
        
        paragrafo_trabalho = doc.add_paragraph(f'{empresa}')
        paragrafo_trabalho.runs[0].bold = True
        doc.add_paragraph(f'Cargo: {cargo}')
        doc.add_paragraph(f'Período: {trabalho.get("periodo", "Período Não Especificado")}')
        
        atividades = trabalho.get('atividades', [])
        if atividades:
            for atividade in atividades:
                texto_atividade = atividade.get('atividade', '') if isinstance(atividade, dict) else str(atividade)
                paragrafo_atividade = doc.add_paragraph(f' • {texto_atividade}')
                paragrafo_atividade.paragraph_format.left_indent = Pt(18)
        
        projetos = trabalho.get('projetos', [])
        if projetos:
            paragrafo_projetos = doc.add_paragraph('Projetos:')
            paragrafo_projetos.runs[0].bold = True
            paragrafo_projetos.paragraph_format.left_indent = Pt(18)
            for projeto in projetos:
                projeto_titulo = projeto.get('titulo', 'Projeto Não Especificado')
                projeto_descricao = projeto.get('descricao', '')
                doc.add_paragraph(f'{projeto_titulo}').paragraph_format.left_indent = Pt(24)
                doc.add_paragraph(projeto_descricao).paragraph_format.left_indent = Pt(24)
        
    # Salvar o documento
    doc.save(arquivo_saida)
    print(f"Currículo salvo em {arquivo_saida}")
def process_text(texto):
    """Processa o texto e retorna JSON estruturado com tratamento de erros aprimorado."""
    chave_api = os.getenv('OPENAI_API_KEY')
    
    if not chave_api:
        raise ValueError("Chave de API OpenAI não encontrada. Defina OPENAI_API_KEY no arquivo .env.")

    modelo_prompt = """
    Você é um especialista em extração de informações estruturadas de currículos. 
    Analise cuidadosamente o texto completo do CV e extraia os detalhes em um formato JSON estruturado e preciso. 
    Priorize a clareza e a completude.

    TEXTO DO CV:
    {texto}
    
    INSTRUÇÕES PARA O OUTPUT:
    - Retorne um JSON completo com todos os campos esperados
    - Use um texto descritivo e conciso
    - Se faltar alguma informação, use strings ou listas vazias
    - Garanta formatação e legibilidade adequadas

    ESTRUTURA DE JSON EXIGIDA:
    {{
        "informacoes_pessoais": {{
            "nome": "Nome Completo",
            "cidade": "Cidade, Estado/País",
            "bairro": "Bairro Opcional",
            "email": "email@exemplo.com",
            "telefone": "Telefone Opcional",
            "cargo": "Cargo Atual ou Desejado"
        }},
        "resumo_qualificacoes": [{{
            "resumo": "Visão geral profissional breve",
            "qualificacoes_chave": [
                {{"qualificacao": "Habilidade ou realização importante"}},
                {{"qualificacao": "Outra habilidade importante"}}
            ]
        }}],
        "experiencia_profissional": [
            {{
                "empresa": "Nome da Empresa",
                "cargo": "Título do Cargo",
                "periodo": "Data de Início - Data de Término",
                "atividades": [
                    {{"atividade": "Responsabilidade ou realização chave"}},
                    {{"atividade": "Outra responsabilidade importante"}}
                ],
                "projetos": [
                    {{"titulo": "Nome do Projeto", "descricao": "Descrição do projeto"}}
                ]
            }}
        ],
        "educacao": [
            {{
                "instituicao": "Nome da Escola/Universidade",
                "grau": "Grau ou Certificação",
                "ano_formatura": "Ano"
            }}
        ],
        "certificacoes": [
            {{"certificado": "Nome da Certificação"}}
        ]
    }}
    """
    
    try:
        prompt = PromptTemplate(template=modelo_prompt, input_variables=["texto"])
        llm = ChatOpenAI(api_key=chave_api, temperature=0, model="gpt-4o-mini")
        resultado = llm.invoke(prompt.format(texto=texto))
        
        print("Resposta Completa do Modelo:")
        print(resultado.content)
        
        dados_json = None
        try:
            dados_json = json.loads(resultado.content)
        except json.JSONDecodeError:
            correspondencia_json = re.search(r'\{.*\}', resultado.content, re.DOTALL | re.MULTILINE)
            if correspondencia_json:
                try:
                    dados_json = json.loads(correspondencia_json.group(0))
                except Exception as e:
                    print(f"Erro de extração JSON: {e}")
        
        estrutura_padrao = {
            "informacoes_pessoais": {
                "nome": "",
                "cidade": "",
                "email": "",
                "telefone": "",
                "cargo": ""
            },
            "resumo_qualificacoes": [],
            "experiencia_profissional": [],
            "educacao": [],
            "certificacoes": []
        }
        
        for chave in estrutura_padrao:
            if chave not in dados_json:
                dados_json[chave] = estrutura_padrao[chave]
        
        return dados_json
    
    except Exception as e:
        print("Erro detalhado ao processar texto:")
        print(traceback.format_exc())
        
        return {
            "informacoes_pessoais": {
                "nome": "",
                "cidade": "",
                "email": "",
                "telefone": "",
                "cargo": ""
            },
            "resumo_qualificacoes": [],
            "experiencia_profissional": [],
            "educacao": [],
            "certificacoes": []
        }

def clear_text(texto):
    """Limpa e normaliza o texto extraído."""
    texto = re.sub(r'\s+', ' ', texto)
    texto = re.sub(r'\n*Página \d+ de \d+\n*', '', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    texto = texto.strip()
    
    return texto

def extract_text_from_pdf(caminho_pdf):
    """Extrai o texto de um arquivo PDF."""
    try:
        leitor = PdfReader(caminho_pdf)
        texto = ""
        for pagina in leitor.pages:
            texto += pagina.extract_text() + "\n"
        
        texto_limpo = clear_text(texto)
        
        return texto_limpo
    
    except Exception as e:
        print(f"Erro ao extrair texto do PDF: {e}")
        return ""

def main(caminho_pdf):
    """Função principal para processar PDF e gerar documento de currículo."""
    texto_pdf = extract_text_from_pdf(caminho_pdf)
    
    if not texto_pdf:
        print("Nenhum texto pôde ser extraído do PDF.")
        return
    
    dados_json = process_text(texto_pdf)
    
    with open('dados_curriculo_extraidos.json', 'w', encoding='utf-8') as f:
        json.dump(dados_json, f, indent=2)
    
    create_docx_from_json('dados_curriculo_extraidos.json')

# Exemplo de uso
if __name__ == "__main__":
    caminho_pdf = "Profile (11).pdf"  # Especifique o seu PDF
    main(caminho_pdf)