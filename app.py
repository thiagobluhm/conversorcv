import os
import streamlit as st
import json
import traceback
import tempfile
from dotenv import load_dotenv
import openai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from pathlib import Path
import base64
import re

os.chdir(os.path.abspath(os.curdir))

def validate_json(dados, estrutura_padrao):
    """Valida e completa o JSON com estrutura padrão."""
    for chave in estrutura_padrao:
        if chave not in dados:
            dados[chave] = estrutura_padrao[chave]
    return dados



def create_docx_from_json(arquivo_json, arquivo_saida='curriculo.docx', logo_path='portfoliologotech.png'):
    """Cria um documento Word formatado a partir de dados de um currículo em JSON e adiciona um logo."""
    try:
        with open(arquivo_json, 'r', encoding='utf-8') as f:
            dados = json.load(f)

        estrutura_padrao = {
            "informacoes_pessoais": {"nome": "", "cidade": "", "email": "", "telefone": "", "cargo": ""},
            "resumo_qualificacoes": [],
            "experiencia_profissional": [],
            "educacao": [],
            "certificacoes": []
        }
        dados = validate_json(dados, estrutura_padrao)

        doc = Document()
        estilo = doc.styles['Normal']
        estilo.font.name = 'Calibri'
        estilo.font.size = Pt(11)
        estilo.font.color.rgb = RGBColor(0, 0, 0)

        def adicionar_espaco():
            """Adiciona um parágrafo vazio para espaçamento."""
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Adicionar logo ao cabeçalho
        if logo_path:
            header = doc.sections[0].header
            header_paragraph = header.paragraphs[0]
            run = header_paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))  # Ajusta o tamanho do logo
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Alinha à direita

        # Informações pessoais
        informacoes_pessoais = dados.get('informacoes_pessoais', {})
        nome = informacoes_pessoais.get('nome', 'Nome Não Encontrado')
        paragrafo_nome = doc.add_paragraph(nome)
        if paragrafo_nome.runs:
            nome_run = paragrafo_nome.runs[0]
            nome_run.bold = True
            nome_run.font.size = Pt(16)
        paragrafo_nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        adicionar_espaco()
        contato = f"Cidade: {informacoes_pessoais.get('cidade', 'N/A')}\nEmail: {informacoes_pessoais.get('email', 'N/A')}\nTelefone: {informacoes_pessoais.get('telefone', 'N/A')}\nPosição: {informacoes_pessoais.get('cargo', 'N/A')}"
        doc.add_paragraph(contato)

        adicionar_espaco()

        # Resumo de qualificações
        doc.add_heading('Resumo de Qualificações', level=2)
        for qualificacao in dados.get('resumo_qualificacoes', []):
            doc.add_paragraph(f"- {qualificacao}")

        adicionar_espaco()

        # Experiência profissional
        doc.add_heading('Experiência Profissional', level=2)
        for experiencia in dados.get('experiencia_profissional', []):
            empresa = experiencia.get('empresa', 'Empresa Não Informada')
            cargo = experiencia.get('cargo', 'Cargo Não Informado')
            periodo = experiencia.get('periodo', 'Período Não Informado')
            local = experiencia.get('local', 'Local Não Informado')
            atividades = experiencia.get('atividades_exercidas', [])

            doc.add_paragraph(f"{empresa} ({local})", style='Heading 3')
            doc.add_paragraph(f"{cargo} - {periodo}", style='Normal')
            
            if atividades:
                doc.add_paragraph("Atividades exercidas:", style='Normal')
                for atividade in atividades:
                    doc.add_paragraph(f"• {atividade}", style='List Bullet')

            ferramentas = experiencia.get('ferramentas', [])
            if ferramentas:
                doc.add_paragraph("Ferramentas utilizadas:", style='Normal')
                for ferramenta in ferramentas:
                    doc.add_paragraph(f"• {ferramenta}", style='List Bullet')

        adicionar_espaco()

        # Educação
        doc.add_heading('Educação', level=2)
        for educacao in dados.get('educacao', []):
            instituicao = educacao.get('instituicao', 'Instituição Não Informada')
            curso = educacao.get('curso', 'Curso Não Informado')
            periodo = educacao.get('periodo', 'Período Não Informado')

            doc.add_paragraph(f"{instituicao}", style='Heading 3')
            doc.add_paragraph(f"{curso} - {periodo}", style='Normal')

        adicionar_espaco()

        # Certificações
        doc.add_heading('Certificações', level=2)
        for certificacao in dados.get('certificacoes', []):
            doc.add_paragraph(f"- {certificacao}", style='Normal')

        # Salvar o documento Word
        doc.save(arquivo_saida)
        print(f"Currículo salvo em {arquivo_saida}")

    except Exception as e:
        print(f"Erro ao criar documento Word: {e}")
        print(traceback.format_exc())




def process_text(texto):
    """Processa o texto e retorna JSON estruturado."""
    load_dotenv()
    chave_api = os.getenv('OPENAI_API_KEY')
    openai.api_key = chave_api

    if not chave_api:
        st.error("Chave da API OpenAI não encontrada.")
        return {}

    modelo_prompt = f"""
                        TEXTO DO CURRÍCULO:
                        {texto}

                        Campos esperados e explicações:
                        1. **informacoes_pessoais**: 
                            Contém as informações pessoais do candidato, incluindo:
                            - "nome": Nome completo do candidato.
                            - "cidade": Cidade e estado de residência.
                            - "email": Endereço de e-mail de contato.
                            - "telefone": Número de telefone para contato.
                            - "cargo": Cargo atual ou pretendido.

                        2. **resumo_qualificacoes**:
                            Lista com as principais habilidades, competências ou realizações do candidato, como:
                            - Conhecimentos técnicos (ex.: Power BI, Python, SQL).
                            - Soft skills (ex.: liderança, trabalho em equipe).
                            - Principais realizações (ex.: "Aumentou a eficiência em X% ao implementar [projeto]").

                        3. **experiencia_profissional**:
                            Lista de experiências profissionais relevantes. Cada entrada deve conter:
                            - "empresa": Nome da empresa.
                            - "cargo": Cargo exercido.
                            - "periodo": Período de atuação (ex.: Janeiro de 2020 - Dezembro de 2022).
                            - "local": Local onde o trabalho foi realizado (ex.: Remoto ou Cidade/Estado).
                            - "atividades_exercidas": Lista de atividades e responsabilidades no cargo. Detalhe as principais contribuições e tarefas realizadas.
                            - "ferramentas": Lista das ferramentas, softwares ou tecnologias utilizadas no cargo (ex.: Power BI, Python, SQL, Tableau).

                        4. **educacao**:
                            Lista de formações acadêmicas do candidato. Cada entrada deve conter:
                            - "instituicao": Nome da instituição de ensino.
                            - "curso": Curso ou programa concluído.
                            - "periodo": Período de realização (ex.: Janeiro de 2016 - Dezembro de 2020).

                        5. **certificacoes**:
                            Lista de certificações relevantes obtidas pelo candidato. Cada entrada deve conter:
                            - Nome da certificação (ex.: "Microsoft Certified: Data Analyst Associate").
                            - Instituição que emitiu a certificação (ex.: Microsoft, AWS, etc.).

                        Formato esperado do JSON de saída:
                        {{
                            "informacoes_pessoais": {{
                                "nome": "",
                                "cidade": "",
                                "email": "",
                                "telefone": "",
                                "cargo": ""
                            }},
                            "resumo_qualificacoes": [
                                "Resumo 1",
                                "Resumo 2"
                            ],
                            "experiencia_profissional": [
                                {{
                                    "empresa": "Empresa X",
                                    "cargo": "Cargo Y",
                                    "periodo": "Janeiro de 2020 - Dezembro de 2022",
                                    "local": "Cidade/Estado",
                                    "atividades_exercidas": [
                                        "Atividade 1",
                                        "Atividade 2"
                                    ],
                                    "ferramentas": [
                                        "Ferramenta 1",
                                        "Ferramenta 2"
                                    ]
                                }}
                            ],
                            "educacao": [
                                {{
                                    "instituicao": "Instituição A",
                                    "curso": "Curso B",
                                    "periodo": "Janeiro de 2016 - Dezembro de 2020"
                                }}
                            ],
                            "certificacoes": [
                                "Certificação 1",
                                "Certificação 2"
                            ]
                        }}
                        """


    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": """Você é um especialista em análise de currículos e extração de informações. 
                                                 Dê sua resposta APENAS com o json solicitado e nada mais. NÃO ESCREVA ```json na resposta!
                """},
                {"role": "user", "content": modelo_prompt}
            ],
            temperature=0,
            max_tokens=4096
        )
          
        conteudo = response.choices[0].message.content.replace("```json", "").strip()
        st.write(f"CONTEUDO: {conteudo}")

        try:
            return json.loads(conteudo)
        except json.JSONDecodeError:
            print("Erro ao converter resposta da API para JSON.")
            return {}
    except Exception as e:
        print(f"Erro ao processar texto com a API OpenAI: {e}")
        return {}

def extract_text_from_pdf(caminho_pdf):
    """Extrai o texto de um arquivo PDF."""
    try:
        leitor = PdfReader(caminho_pdf)
        texto = "".join(pagina.extract_text() for pagina in leitor.pages)
        return clear_text(texto)
    except Exception as e:
        print(f"Erro ao extrair texto do PDF: {e}")
        return ""

def clear_text(texto):
    """Limpa e normaliza o texto extraído."""
    texto = re.sub(r'\s+', ' ', texto)
    texto = re.sub(r'\n*Página \d+ de \d+\n*', '', texto)
    return texto.strip()

# Função para adicionar uma imagem de fundo a partir de um arquivo local
def add_bg_from_local(image_file):
    with Path(image_file).open("rb") as file:
        encoded_string = base64.b64encode(file.read()).decode()
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-color: #FFFFFF;
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            border-color: rgba(31,216,135,1) ;
 
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

# Função para exibir uma imagem de logo no topo a partir de um arquivo local
def add_logo_from_local(logo_file):
    with Path(logo_file).open("rb") as file:
        encoded_string = base64.b64encode(file.read()).decode()
    st.markdown(
        f"""
        <style>
        [data-testid="stAppViewContainer"] > .main {{
            padding-top: 0px;
        }}
        .logo-container {{
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 1rem 0;
        }}
        .logo-container img {{
            max-width: 200px;
            height: auto;
        }}
        </style>
        <div class="logo-container">
            <img src="data:image/png;base64,{encoded_string}" alt="Logo">
        </div>
        """,
        unsafe_allow_html=True
    )


def main():
    st.set_page_config(page_title="Conversor de CV PDF para DOCX **", page_icon="📄", layout="centered")
    
    #add_bg_from_local("bg.png")
    add_logo_from_local("portfoliologotech.png")

    st.markdown("<h1 style='text-align: center;'>Conversor de Currículo</h1>", unsafe_allow_html=True)

    with st.form(key="upload_form"):
        uploaded_file = st.file_uploader("Envie seu currículo em PDF", type="pdf")
        submit_button = st.form_submit_button("Converter currículo")

    if submit_button and uploaded_file:
        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                temp_pdf.write(uploaded_file.getvalue())
                temp_pdf_path = temp_pdf.name

            status_text.text("Etapa 1: Extraindo texto do PDF...")
            progress_bar.progress(20)
            pdf_text = extract_text_from_pdf(temp_pdf_path)

            if not pdf_text.strip():
                st.error("Não foi possível extrair texto do PDF.")
                return

            #st.write("Texto extraído do PDF:", pdf_text)

            status_text.text("Etapa 2: Processando o texto do currículo...")
            progress_bar.progress(50)
            json_data = process_text(pdf_text)

            # IMPRIMINDO NA TELA O TEXTO EXTRAIDO
            #st.write(json_data)

            if not json_data:
                st.error("Erro ao gerar JSON do currículo.")
                return

            with tempfile.NamedTemporaryFile(delete=False, suffix=".json", mode='w', encoding='utf-8') as temp_json:
                json.dump(json_data, temp_json, indent=2)
                temp_json_path = temp_json.name

            status_text.text("Etapa 3: Convertendo texto para formato Word...")
            progress_bar.progress(80)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                create_docx_from_json(temp_json_path, temp_docx.name)
                temp_docx_path = temp_docx.name

            progress_bar.progress(100)
            st.success("Conversão concluída com sucesso! Baixe seu currículo abaixo.")
            with open(temp_docx_path, "rb") as file:
                st.download_button(
                    label="Baixar currículo em DOCX",
                    data=file.read(),
                    file_name="Curriculo_Formatado.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Ocorreu um erro: {e}")
            st.error(traceback.format_exc())

if __name__ == "__main__":
    main()
